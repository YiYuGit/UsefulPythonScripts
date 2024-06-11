# Read photos in a folder, sort them by creation date, and write them into one docx file to make an album.

import os
import re
from datetime import datetime
from docx import Document
from docx.shared import Inches
from PIL import Image
from io import BytesIO

# Path to the folder containing photos
photos_folder = r'C:\capture'

# Create a list of photo filenames along with their creation dates
photo_files = os.listdir(photos_folder)
photo_info = []
for filename in photo_files:
    if filename.lower().endswith('.png') and not filename.lower().endswith('png.meta'):
        match = re.match(r'(\d{4}-\d{2}-\d{2}_\d{2}-\d{2}-\d{2})\.png', filename)
        if match:
            date_str = match.group(1)
            date_obj = datetime.strptime(date_str, '%Y-%m-%d_%H-%M-%S')
            photo_info.append((filename, date_obj))

# Sort photos by creation date
photo_info.sort(key=lambda x: x[1])

# Create a new Word document
doc = Document()
section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
left_margin = Inches(1)
right_margin = Inches(1)
top_margin = Inches(1)
bottom_margin = Inches(1)
section.left_margin = left_margin
section.right_margin = right_margin
section.top_margin = top_margin
section.bottom_margin = bottom_margin

# Add photos to the document
photos_per_page = 2
for i in range(0, len(photo_info), photos_per_page):
    if i > 0:
        doc.add_page_break()
    for j in range(photos_per_page):
        idx = i + j
        if idx < len(photo_info):
            photo_filename, _ = photo_info[idx]
            photo_path = os.path.join(photos_folder, photo_filename)
            img = Image.open(photo_path)
            img.thumbnail((Inches(6), Inches(6)))
            doc.add_picture(photo_path, width=Inches(6))
            doc.add_paragraph(photo_filename)

# Save the document
# doc.save('photo_album_without_meta.docx')


# Another version. Save the document in the same folder as the photos
docx_filename = os.path.join(photos_folder, 'photo_album.docx')
doc.save(docx_filename)
print(f"Document saved as '{docx_filename}'")