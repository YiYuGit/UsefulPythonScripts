import os
import re
import sys
import time
from datetime import datetime
from docx import Document
from docx.shared import Inches
from docx.shared import Pt
from io import BytesIO
from docx.enum.text import WD_ALIGN_PARAGRAPH
from PIL import Image as PILImage
from exif import Image as ExifImage


"""
Overall, this script is designed to read information from txt and jpg file, write info from txt to jpg and save as jpg in a new folder.
Then, docx will read the jpgs and use some information from txt to generate a word file report


Older comments:
In this example, the script can read all photo in a folder, and make then into a photo album. The script also read another txt file folder and use file name to find the matching txt file for each photo.
In this script. The code:
1. find all photos by the name format.
2. format the docx file.
3. two loop for inserting two photos per page
4. make 3by3 table for each photo, merge cells as needed.
5. insert photo
6. search for the txt file with the same file name
7. read the 3 lines in the txt file
8. insert the texts into the desired cells
9. same the fine into a .docx file

Can be built into exe using pyinstaller. The file path issue is a problem. Currently, the absolute path works.
"""



def decimal_to_dms(decimal_degree):
    degrees = int(decimal_degree)
    minutes_decimal = (decimal_degree - degrees) * 60
    minutes = int(minutes_decimal)
    seconds = (minutes_decimal - minutes) * 60
    return degrees, minutes, seconds

def read_and_convert_coordinates(file_path):
    with open(file_path, 'r') as file:
        # Read latitude and longitude from the file
        decimal_latitude = float(file.readline().strip())
        decimal_longitude = float(file.readline().strip())
        altitude_in_meter = float(file.readline().strip())
        timestamp = str(file.readline().strip())
        direction = str(file.readline().strip())
        no_use = str(file.readline().strip())
        description = str(file.readline().strip())
        
        direction = direction[30:33]
        direction_numberic_part = ''.join(filter(str.isdigit, direction))
        direction = direction_numberic_part
        

    # Convert latitude to degrees, minutes, and decimal seconds
    dms_lat = decimal_to_dms(abs(decimal_latitude))
    lat_ref = 'N' if decimal_latitude >= 0 else 'S'

    # Convert longitude to degrees, minutes, and decimal seconds
    dms_long = decimal_to_dms(abs(decimal_longitude))
    long_ref = 'E' if decimal_longitude >= 0 else 'W'

    return dms_lat, lat_ref, dms_long, long_ref, altitude_in_meter, timestamp, direction, description






# Path to the folder containing photos, 'r' means raw string, Prefacing the string definition with ‘r’ is a useful way to 
# define a string where you need the backslash to be an actual backslash and not part of an escape code that means something else in the string.
#photos_folder = r'Y:\NewPhotoLog\png'
#text_folder = r'Y:\NewPhotoLog\txt'
#photos_folder = os.getcwd()

#absolute_path = os.path.dirname(__file__)
absolute_path = os.path.dirname(os.path.abspath(__file__))
#absolute_path = os.path.dirname(sys.executable)

# Command to generate the pysintaller is "pyinstaller --onefile WriteToJpgExifAndGenPhotoLog.py"



relative_photo_path = "capture"
relative_txt_path = "Txt_log"
relative_new_photo_path = "new_photo"

photos_folder = os.path.join(absolute_path, relative_photo_path)
text_folder = os.path.join(absolute_path,relative_txt_path)
new_photo_folder = os.path.join(absolute_path, relative_new_photo_path)
doc = Document(os.path.join(absolute_path, "E1527-21.docx"))



# Create a list of photo filenames along with their creation dates
photo_files = os.listdir(photos_folder)
text_files = os.listdir(text_folder)



# List all JPG file in the folder
jpg_files = [file for file in photo_files if file.endswith('.jpg') and not file.lower().endswith('jpg.meta')]

for jpg_file in jpg_files:
    jpg_path = os.path.join(photos_folder, jpg_file)
    
    text_filename = os.path.splitext(jpg_file)[0] + '.txt'
    text_path = os.path.join(text_folder, text_filename)
    dms_lat, lat_ref, dms_long, long_ref, altitude_in_meter, timestamp, direction, description = read_and_convert_coordinates(text_path)
    
    with open(jpg_path, 'rb') as image_file:
        my_image = ExifImage(image_file)
        
    my_image.make = "Your Camera Make"
    my_image.model = "Your Camera Model"
    my_image.gps_latitude_ref = lat_ref
    my_image.gps_longitude_ref = long_ref
    my_image.gps_latitude = dms_lat
    my_image.gps_longitude = dms_long
    my_image.gps_altitude = altitude_in_meter #in meters
    my_image.gps_altitude_ref = 0  # 0 means above sea level
    #my_image.gps_timestamp = timestamp
    my_image.gps_datestamp = timestamp[0:10]
    gps_timestamp = timestamp[0:4] +":" + timestamp[5:7] +":" + timestamp[8:10] + " " +timestamp[11:13] + ":"+ timestamp[14:16] + ":"+ timestamp[17:19]
    my_image.datetime = gps_timestamp
    my_image.datetime_original = gps_timestamp
    my_image.gps_img_direction = direction
    my_image.gps_img_direction_ref = "T"
    my_image.user_comment = description
    
    # for reference 2023-11-16T16:28:01Z
    new_photo_name = os.path.splitext(jpg_file)[0] + '.jpg'
    new_photo_path = os.path.join(new_photo_folder, new_photo_name)
    with open(new_photo_path, 'wb') as new_photo_file:
        new_photo_file.write(my_image.get_file()) 





# After writing info to new photos 
new_photo_files = os.listdir(new_photo_folder)
text_files = os.listdir(text_folder)



photo_info = []
for filename in new_photo_files:
    if filename.lower().endswith('.jpg') and not filename.lower().endswith('jpg.meta'):
        match = re.match(r'(\d{4}-\d{2}-\d{2}_\d{2}-\d{2}-\d{2})\.jpg', filename)
        if match:
            date_str = match.group(1)
            date_obj = datetime.strptime(date_str, '%Y-%m-%d_%H-%M-%S')
            photo_info.append((filename, date_obj))

# Sort photos by creation date
photo_info.sort(key=lambda x: x[1])

# Depends on the acutal need, user may create a new Word document or use existing file template 
# If using file template, make sure the template file already set the table style to "Table Grid", or there may be error.

# Create a new Word document
#doc = Document()

# Using the template, set the template file path
#doc = Document(r'Y:\capture\E1527-21.docx')
#doc = Document(os.path.join(sys.path[0], "E1527-21.docx"))

# Set the page size, margin, etc. Make sure all needed are in the "from ... import ..."
section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
left_margin = Inches(0.5)
right_margin = Inches(0.5)
top_margin = Inches(0.4)
bottom_margin = Inches(0.4)
section.left_margin = left_margin
section.right_margin = right_margin
section.top_margin = top_margin
section.bottom_margin = bottom_margin

#doc.add_paragraph('    ')


# Add photos to the document, photos_per_page is used to split the pages for every two photo
photos_per_page = 2
for i in range(0, len(photo_info), photos_per_page):
    if i > 0:
        doc.add_page_break()
        
    for j in range(photos_per_page):
        idx = i + j
        if idx < len(photo_info):
            # Read the photo file name from photo_info, throw out useless part
            photo_filename, _ = photo_info[idx]
            photo_path = os.path.join(photos_folder, photo_filename)
            img = PILImage.open(photo_path)
            # Make this image into a thumbnail. This method modifies the image to contain a thumbnail version of itself, no larger than the given size. Not needed for writing in word file
            #img.thumbnail((Inches(6), Inches(6)))
            
            # Create tables and merge certain cells to hold photo. Refer to python docx documentation for how to merge cells
            table = doc.add_table(3,3, style='Table Grid')

            # Can use autofit for the table. In this case, we have fixed width from the template, read the numbers from template and use them here.
            #table.allow_autofit = True
            table.columns[0].width = Inches(0.84)
            table.columns[1].width = Inches(0.84)
            table.columns[1].width = Inches(5.83)

            # Table Grid can be changed during the creation of the table, see previous lines
            #table.style = 'Table Grid'
            
            # The initial table is 3 by 3, to match the required format, need to merge some cells
            a = table.cell(1,0)
            b = table.cell(1,1)
            A = a.merge(b)
            #A.text = "Direction Photo Taken:" + "\n"+ "\n"
            
            # Merge some other cells
            c = table.cell(2,0)
            d = table.cell(2,1)
            B = c.merge(d)
            
            # Read the corresponding text file,  " text_lines = text_file.read().splitlines()[:7]" means select first 7 lines
            text_filename = os.path.splitext(photo_filename)[0] + '.txt'
            text_path = os.path.join(text_folder, text_filename)
            with open(text_path, 'r') as text_file:
                text_lines = text_file.read().splitlines()[:7]   
                
            # Populate the table cells with text from the txt file. The first is an example for writing one line in txt into a cell in the table. 
            # The second line is an example for writing two line in txt into the same cell
            A.text = text_lines[4]
            B.text = f"{text_lines[5]} {text_lines[6]}"
                    
            
            # Get date info from txt file line 4(timestamp in the format of (2023-11-16T16:28:01Z) and reformat into mm/dd/yyyy.  
            #date = photo_filename[5:7]+"/"+photo_filename[8:10]+"/"+photo_filename[0:4]
            timestamp_text = text_lines[3]
            date = timestamp_text[5:7]+"/"+timestamp_text[8:10]+"/"+timestamp_text[0:4]

            # Write the date to cell, also, change the font size to 8, based on template requirement
            table.cell(0,0).text = "Date" + "\n" + date
            table.cell(0,0).paragraphs[0].runs[0].font.size = Pt(8)

            # Do the same for the Photo numbner, the index start from 1, change font size to 8
            table.cell(0,1).text = "Photo"+ "\n"  "No." + str(idx +1)
            table.cell(0,1).paragraphs[0].runs[0].font.size = Pt(8)

            # Merge the cell for the photo cell
            e = table.cell(0,2)
            f = table.cell(2,2)
            C = e.merge(f)

            # To match the format, add empty run in the photo cell before inserting the photo. The photo width is set to 6 inches.
            paragraph = table.cell(1,2).paragraphs[0]
            paragraph.add_run("\n")
            paragraph.add_run().add_picture(photo_path, width = Inches(6))
            paragraph.add_run("\n")
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            
            
            
            # Add one empty paragrapg to separate the two tables on the same page
            doc.add_paragraph('    ')


# After writing the file, set the font for the whole file, in this case, used 'Univers LT 45 Light' and Bold 
font =  doc.styles['Normal'].font
font.name = 'Univers LT 45 Light'
font.bold = True

# Save the document in the same folder as the photos
docx_filename = os.path.join(new_photo_folder, 'PhotoLogAll.docx')
doc.save(docx_filename)

# Sleep for 4 seconds, not necessary, just to be safe
time.sleep(4)

# Open the Word document
os.startfile(docx_filename)

# Open the folder of the new_photo 
os.startfile(new_photo_folder)
print(f"Document saved as '{docx_filename}'")