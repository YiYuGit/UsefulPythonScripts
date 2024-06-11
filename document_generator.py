# Generate a docx file with image and text


from docx import Document
from docx.shared import Inches
import sys
import os

def generate_document(text_file, image_file, output_path):
    # Read text from file
    with open(text_file, 'r') as f:
        text = f.read()

    # Create a new Word document
    doc = Document()
    doc.add_paragraph(text)

    # Add the image to the document
    doc.add_picture(image_file, width=Inches(4))  # Adjust width as needed

    # Save the document with the specified output path and filename
    output_filename = os.path.join(output_path, 'output.docx')
    doc.save(output_filename)

if __name__ == "__main__":
    if len(sys.argv) != 4:
        print("Usage: python document_generator.py <text_file_path> <image_file_path> <output_directory>")
        sys.exit(1)

    text_file_path = sys.argv[1]
    image_file_path = sys.argv[2]
    output_path = sys.argv[3]

    generate_document(text_file_path, image_file_path, output_path)





#Example

#python document_generator.py "C:\xxxxxx\xxxxx.txt" "C:\xxxxx\xxxx.jpg" "C:\xxxxx\xxxxx\xxxx" 